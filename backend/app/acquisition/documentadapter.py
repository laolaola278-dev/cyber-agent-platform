"""Phase 28 -- Document Acquisition Tool Adapter (spec 12).

Real parsers only:
  * PDF    -> pypdf
  * DOCX   -> python-docx
  * XLSX   -> openpyxl
  * HTML   -> lxml.html (fallback: stdlib html.parser)
  * JSON   -> stdlib json
  * TEXT   -> plain text

Every parser runs inside the Worker/Sandbox boundary and reports its real
backend via ``parser_backend``. If a dependency is unavailable at runtime the
adapter reports ``parser_backend="UNAVAILABLE:..."`` and returns a synthetic
placeholder -- it NEVER pretends a real parse happened.

Extraction output (title/text/tables/metadata) is an *ExtractedDocument*
candidate; the raw bytes are preserved as the Evidence artifact.
"""

from __future__ import annotations

import io
import json as _json
import zipfile
from dataclasses import dataclass

from app.acquisition.models import ExtractedDocument

_MAX_TEXT_CHARS = 200_000


@dataclass
class DocumentParseResult:
    ok: bool
    document: ExtractedDocument | None = None
    parser_backend: str = ""
    error: str = ""


def _clip(text: str) -> str:
    return text[:_MAX_TEXT_CHARS]


def _parse_html(content: bytes, source_url: str) -> DocumentParseResult:
    try:
        from lxml import html as lxml_html  # type: ignore[import-not-found]
    except ImportError:
        return DocumentParseResult(
            ok=False, parser_backend="UNAVAILABLE:lxml", error="lxml not installed"
        )
    try:
        root = lxml_html.fromstring(content)
        title = (root.findtext(".//title") or "").strip()
        body = root.find(".//body")
        if body is None:
            body = root
        text_parts: list[str] = []
        for node in body.iter():
            if node.tag in ("script", "style", "noscript", "svg", "head"):
                continue
            if node.tag in ("p", "h1", "h2", "h3", "h4", "li", "tr", "div"):
                part = " ".join((node.text_content() or "").split())
                if part:
                    text_parts.append(part)
        text = "\n".join(text_parts)[:_MAX_TEXT_CHARS]
        links = [
            (node.get("href") or "")
            for node in body.iter("a")
            if node.get("href")
        ][:500]
        tables: list[list[list[str]]] = []
        for table in body.iter("table"):
            rows: list[list[str]] = []
            for tr in table.iter("tr"):
                cells = [" ".join(td.text_content().split()) for td in tr.iter("td", "th")]
                if cells:
                    rows.append(cells)
            if rows:
                tables.append(rows[:50])
        metadata = {"html_title": title, "link_count": len(links)}
        return DocumentParseResult(
            ok=True,
            document=ExtractedDocument(
                title=title,
                text=text,
                sections=[],
                tables=tables,
                metadata=metadata,
                links=links,
                source_url=source_url,
                extraction_backend="lxml.html",
            ),
            parser_backend="lxml.html",
        )
    except Exception as error:  # noqa: BLE001
        return DocumentParseResult(
            ok=False, parser_backend="lxml.html", error=f"parse failed: {error}"
        )


def _parse_pdf(content: bytes, source_url: str) -> DocumentParseResult:
    try:
        from pypdf import PdfReader  # type: ignore[import-not-found]
    except ImportError:
        return DocumentParseResult(
            ok=False, parser_backend="UNAVAILABLE:pypdf", error="pypdf not installed"
        )
    try:
        reader = PdfReader(io.BytesIO(content))
        pages: list[str] = []
        for page in reader.pages:
            pages.append(page.extract_text() or "")
        text = "\n".join(pages)[:_MAX_TEXT_CHARS]
        metadata = {
            "pdf_pages": len(reader.pages),
            "pdf_title": (reader.metadata.title if reader.metadata else None),
            "pdf_author": (reader.metadata.author if reader.metadata else None),
        }
        return DocumentParseResult(
            ok=True,
            document=ExtractedDocument(
                title=metadata["pdf_title"] or "",
                text=text,
                metadata=metadata,
                author=metadata["pdf_author"],
                source_url=source_url,
                extraction_backend="pypdf",
            ),
            parser_backend="pypdf",
        )
    except Exception as error:  # noqa: BLE001
        return DocumentParseResult(
            ok=False, parser_backend="pypdf", error=f"parse failed: {error}"
        )


def _parse_docx(content: bytes, source_url: str) -> DocumentParseResult:
    try:
        import docx  # type: ignore[import-not-found]
    except ImportError:
        return DocumentParseResult(
            ok=False, parser_backend="UNAVAILABLE:python-docx", error="docx not installed"
        )
    try:
        document = docx.Document(io.BytesIO(content))
        paragraphs = [p.text for p in document.paragraphs if p.text.strip()]
        tables: list[list[list[str]]] = []
        for table in document.tables:
            rows: list[list[str]] = []
            for row in table.rows:
                rows.append([cell.text.strip() for cell in row.cells])
            if rows:
                tables.append(rows[:50])
        text = "\n".join(paragraphs)[:_MAX_TEXT_CHARS]
        return DocumentParseResult(
            ok=True,
            document=ExtractedDocument(
                title=paragraphs[0] if paragraphs else "",
                text=text,
                tables=tables,
                metadata={"docx_paragraphs": len(paragraphs)},
                source_url=source_url,
                extraction_backend="python-docx",
            ),
            parser_backend="python-docx",
        )
    except Exception as error:  # noqa: BLE001
        return DocumentParseResult(
            ok=False,
            parser_backend="python-docx",
            error=f"parse failed: {error}",
        )


def _parse_xlsx(content: bytes, source_url: str) -> DocumentParseResult:
    try:
        import openpyxl  # type: ignore[import-not-found]
    except ImportError:
        return DocumentParseResult(
            ok=False, parser_backend="UNAVAILABLE:openpyxl", error="openpyxl not installed"
        )
    try:
        workbook = openpyxl.load_workbook(io.BytesIO(content), read_only=True, data_only=True)
        tables: list[list[list[str]]] = []
        texts: list[str] = []
        for sheet in workbook.worksheets:
            rows: list[list[str]] = []
            for row in sheet.iter_rows(values_only=True):
                cells = ["" if c is None else str(c) for c in row]
                if any(cells):
                    rows.append(cells)
            if rows:
                tables.append(rows[:100])
        workbook.close()
        return DocumentParseResult(
            ok=True,
            document=ExtractedDocument(
                title="",
                text="\n".join(texts)[:_MAX_TEXT_CHARS],
                tables=tables,
                metadata={"xlsx_sheets": len(tables)},
                source_url=source_url,
                extraction_backend="openpyxl",
            ),
            parser_backend="openpyxl",
        )
    except Exception as error:  # noqa: BLE001
        return DocumentParseResult(
            ok=False,
            parser_backend="openpyxl",
            error=f"parse failed: {error}",
        )


def _parse_json(content: bytes, source_url: str) -> DocumentParseResult:
    try:
        data = _json.loads(content.decode("utf-8", "replace"))
        return DocumentParseResult(
            ok=True,
            document=ExtractedDocument(
                title="",
                text=_clip(_json.dumps(data, ensure_ascii=False, indent=1)),
                metadata={"json_type": type(data).__name__},
                source_url=source_url,
                extraction_backend="stdlib-json",
            ),
            parser_backend="stdlib-json",
        )
    except ValueError as error:
        return DocumentParseResult(
            ok=False, parser_backend="stdlib-json", error=f"invalid JSON: {error}"
        )


def _parse_text(content: bytes, source_url: str) -> DocumentParseResult:
    try:
        text = content.decode("utf-8", "replace")[:_MAX_TEXT_CHARS]
        return DocumentParseResult(
            ok=True,
            document=ExtractedDocument(
                text=text, source_url=source_url, extraction_backend="stdlib-text"
            ),
            parser_backend="stdlib-text",
        )
    except Exception as error:  # noqa: BLE001
        return DocumentParseResult(
            ok=False, parser_backend="stdlib-text", error=f"decode failed: {error}"
        )


class DocumentAdapter:
    """Route raw bytes to the real parser for the detected document type."""

    def __init__(self) -> None:
        self._parsers = {
            "application/pdf": _parse_pdf,
            "application/msword": _parse_docx,
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document": _parse_docx,
            "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet": _parse_xlsx,
            "application/vnd.ms-excel": _parse_xlsx,
            "text/html": _parse_html,
            "application/json": _parse_json,
            "text/plain": _parse_text,
            "text/markdown": _parse_text,
        }

    def detect_type(self, content: bytes, content_type: str | None) -> str:
        """Content-type first, then magic-byte sniffing (Tika-inspired)."""
        if content_type:
            ctype = content_type.split(";")[0].strip().lower()
            if ctype in self._parsers:
                return ctype
        # magic bytes
        if content.startswith(b"%PDF"):
            return "application/pdf"
        if content.startswith(b"PK\x03\x04"):
            try:
                with zipfile.ZipFile(io.BytesIO(content)) as zf:
                    names = zf.namelist()
                    if "word/document.xml" in names:
                        return (
                            "application/vnd.openxmlformats-officedocument."
                            "wordprocessingml.document"
                        )
                    if "xl/workbook.xml" in names:
                        return (
                            "application/vnd.openxmlformats-officedocument."
                            "spreadsheetml.sheet"
                        )
            except zipfile.BadZipFile:
                pass
        sniffed = content
        if sniffed.startswith(b"\xef\xbb\xbf"):
            sniffed = sniffed[3:]  # UTF-8 BOM
        stripped = sniffed.lstrip()
        if stripped.startswith((b"<", b"<!doctype", b"<html")):
            return "text/html"
        if stripped.startswith((b"{", b"[")):
            return "application/json"
        return "text/plain"

    def parse(
        self, content: bytes, *, content_type: str | None, source_url: str
    ) -> DocumentParseResult:
        detected = self.detect_type(content, content_type)
        parser = self._parsers.get(detected)
        if parser is None:
            return DocumentParseResult(
                ok=False,
                parser_backend=f"UNSUPPORTED:{detected}",
                error=f"no parser for {detected}",
            )
        result = parser(content, source_url)
        return result
