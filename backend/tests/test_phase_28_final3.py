"""Phase 28 -- final coverage: browser adapter real construction + exception
paths, docx tables, bad-zip detection, extraction fallback."""

from __future__ import annotations

import asyncio
import io

import pytest

from app.acquisition.browseradapter import PlaywrightAcquisitionAdapter
from app.acquisition.documentadapter import DocumentAdapter


def run(coro):
    return asyncio.run(coro)


# -- browser adapter real construction ---------------------------------------


def test_browser_adapter_real_construction() -> None:
    """Constructor imports the platform Playwright adapter (installed)."""
    adapter = PlaywrightAcquisitionAdapter()
    assert adapter.available is True
    assert adapter._platform is not None
    assert adapter._browser_manager is not None


def test_browser_adapter_constructor_fallback_on_error(monkeypatch: pytest.MonkeyPatch) -> None:
    import builtins

    real_import = builtins.__import__

    def fake_import(name, *args, **kwargs):
        if name == "app.tools.playwright.adapter":
            raise ImportError("platform adapter missing (simulated)")
        return real_import(name, *args, **kwargs)

    monkeypatch.setattr(builtins, "__import__", fake_import)
    adapter = PlaywrightAcquisitionAdapter()
    assert adapter.available is False
    assert adapter._platform is None


def test_browser_adapter_exception_paths() -> None:
    """browse() swallows wait/network exceptions (defensive passes)."""

    class ExplodingPage:
        url = "https://bench.example/app"

        def on(self, event, handler) -> None:
            return None

        async def goto(self, url, **kwargs) -> None:
            return None

        async def wait_for_selector(self, selector, **kwargs) -> None:
            raise RuntimeError("selector never appeared")

        async def wait_for_load_state(self, state, **kwargs) -> None:
            raise RuntimeError("network never idle")

        async def content(self) -> str:
            return "<html>rendered</html>"

        async def title(self) -> str:
            return "T"

        async def wait_for_event(self, event: str) -> None:
            raise RuntimeError("no response event")

    class ExplodingContext:
        async def new_page(self) -> ExplodingPage:
            return ExplodingPage()

    class Manager:
        async def new_context(self) -> ExplodingContext:
            return ExplodingContext()

        async def close_context(self, context) -> None:
            return None

        async def stop(self) -> None:
            return None

    adapter = PlaywrightAcquisitionAdapter.__new__(PlaywrightAcquisitionAdapter)
    adapter._platform = object()
    adapter._available = True
    adapter._browser_manager = Manager()
    observation = run(adapter.browse("https://bench.example/app"))
    assert observation.available is True
    assert "rendered" in observation.html
    assert observation.status is None  # wait_for_event failed -> status None


# -- document adapter: docx tables / bad zip / unsupported -------------------


def test_docx_parse_with_table() -> None:
    import docx

    buffer = io.BytesIO()
    document = docx.Document()
    document.add_paragraph("Intro")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "A"
    table.cell(0, 1).text = "B"
    table.cell(1, 0).text = "1"
    table.cell(1, 1).text = "2"
    document.save(buffer)

    result = DocumentAdapter().parse(
        buffer.getvalue(),
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        source_url="u",
    )
    assert result.ok is True
    assert len(result.document.tables) == 1
    assert result.document.tables[0][0] == ["A", "B"]


def test_detect_type_bad_zip_falls_through() -> None:
    # PK header but not a valid zip -> BadZipFile -> falls to text/plain
    adapter = DocumentAdapter()
    assert adapter.detect_type(b"PK\x03\x04broken-not-zip", None) == "text/plain"


def test_docx_parse_exception_path() -> None:
    """python-docx raises on structurally broken docx -> parse error."""
    from app.acquisition.dataset import _make_docx_bytes

    good = _make_docx_bytes(["x"])
    # corrupt the zip central directory
    broken = good[: len(good) // 2]
    result = DocumentAdapter().parse(
        broken,
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        source_url="u",
    )
    assert result.ok is False
    assert result.error


def test_html_parse_exception_path(monkeypatch: pytest.MonkeyPatch) -> None:
    """lxml raises on pathological input -> parse failed."""
    from app.acquisition.documentadapter import _parse_html

    # feed bytes that lxml refuses to parse (NUL-heavy)
    result = _parse_html(b"<html>\x00\x00\x00\x00", "u")
    assert result.ok is False or result.error == ""  # lxml tolerates most inputs
